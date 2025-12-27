import streamlit as st
from docx import Document
from docx.shared import Cm, Pt
from io import BytesIO

# 1. Define JSSD Abbreviation Dictionary (Extracted from Vol II)
ABBR_DICT = {
    "Lieutenant Colonel": "Lt Col",
    "Commanding Officer": "CO",
    "Brigadier": "Brig",
    "dated": "dt",
    "Reference": "Ref"
}

def create_jssd_doc(subject, points, security):
    doc = Document()
    # Apply JSSD 2025 Margins 
    section = doc.sections[0]
    section.left_margin, section.right_margin = Cm(3.5), Cm(2.0)
    section.top_margin, section.bottom_margin = Cm(2.0), Cm(2.0)

    # Security Grading (Top & Bottom, No Underline) [cite: 5064, 5068]
    def add_sec(text):
        p = doc.add_paragraph()
        p.alignment = 1 # Center
        run = p.add_run(text.upper())
        run.bold = True

    add_sec(security)
    
    # Bold Caps Subject [cite: 404, 5056]
    doc.add_paragraph()
    subj = doc.add_paragraph()
    subj.alignment = 1
    run = subj.add_run(f"SUBJECT: {subject.upper()}")
    run.bold = True

    # 1.1.1. Numbering Logic 
    for i, line in enumerate(points.split('\n'), 1):
        if line.strip():
            # Apply Abbreviation Replacer
            for word, abbr in ABBR_DICT.items():
                line = line.replace(word, abbr)
            
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            num = p.add_run(f"{i}. ")
            num.bold = True
            p.add_run(line.strip())

    add_sec(security)
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# Streamlit UI
st.title("JSSD 2025 Service Writing App")
subj_input = st.text_input("Subject")
sec_input = st.selectbox("Security", ["UNCLASSIFIED", "RESTRICTED", "CONFIDENTIAL", "SECRET"])
body_input = st.text_area("Points (One per line)")

if st.button("Download Word File"):
    file_data = create_jssd_doc(subj_input, body_input, sec_input)
    st.download_button("Download .docx", file_data, "JSSD_Doc.docx")
